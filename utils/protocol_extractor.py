"""
프로토콜 추출기 - 논문 및 문서에서 실험 프로토콜 자동 추출

이 모듈은 다양한 형식의 문서에서 실험 프로토콜을 자동으로 추출하고
구조화된 JSON 형식으로 변환합니다.

주요 기능:
1. 다양한 파일 형식 지원 (PDF, Word, HTML, TXT 등)
2. 실험 섹션 자동 식별
3. 재료, 조건, 절차 추출 및 구조화
4. 단위 표준화 및 변환
5. 화학물질/장비 인식
6. 다국어 지원
"""

import os
import re
import json
import time
import hashlib
import logging
import tempfile
import traceback
from pathlib import Path
from typing import Dict, List, Optional, Union, Tuple, Any
from datetime import datetime
from dataclasses import dataclass, field, asdict
from enum import Enum
import threading
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError

# 파일 처리
import PyPDF2
import pdfplumber
from docx import Document
import chardet
from bs4 import BeautifulSoup
import markdown
from striprtf.striprtf import rtf_to_text

# OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# NLP
import spacy
from collections import defaultdict
import unicodedata

# 데이터 처리
import pandas as pd
import numpy as np

# 화학 관련
try:
    from chemparse import parse_formula
    CHEMPARSE_AVAILABLE = True
except ImportError:
    CHEMPARSE_AVAILABLE = False

# 로컬 임포트
from config.error_config import ERROR_MESSAGES
from utils.error_handler import ErrorHandler, error_handler, ErrorContext


# ==================== 데이터 모델 ====================

class DocumentType(Enum):
    """지원되는 문서 타입"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MD = "markdown"
    RTF = "rtf"
    
    @classmethod
    def from_extension(cls, ext: str) -> Optional['DocumentType']:
        """파일 확장자로부터 문서 타입 결정"""
        ext = ext.lower().lstrip('.')
        mapping = {
            'pdf': cls.PDF,
            'docx': cls.DOCX,
            'doc': cls.DOCX,
            'txt': cls.TXT,
            'html': cls.HTML,
            'htm': cls.HTML,
            'md': cls.MD,
            'markdown': cls.MD,
            'rtf': cls.RTF
        }
        return mapping.get(ext)


@dataclass
class Material:
    """재료 정보"""
    name: str
    cas_number: Optional[str] = None
    supplier: Optional[str] = None
    purity: Optional[str] = None
    grade: Optional[str] = None
    amount: Optional[str] = None
    unit: Optional[str] = None
    concentration: Optional[str] = None  # 추가
    concentration_unit: Optional[str] = None  # 추가
    category: str = "chemical"
    synonyms: List[str] = field(default_factory=list)
    properties: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Equipment:
    """장비 정보"""
    name: str
    model: Optional[str] = None
    manufacturer: Optional[str] = None
    settings: Dict[str, Any] = field(default_factory=dict)
    calibration: Optional[Dict] = None


@dataclass
class Condition:
    """실험 조건"""
    parameter: str
    value: Union[float, str]
    unit: Optional[str] = None
    tolerance: Optional[float] = None
    profile: Optional[List[Dict]] = None  # 시간-값 프로파일


@dataclass
class Step:
    """실험 단계"""
    number: int
    action: str
    duration: Optional[str] = None
    temperature: Optional[float] = None
    conditions: List[Condition] = field(default_factory=list)
    notes: Optional[str] = None
    critical: bool = False


@dataclass
class Protocol:
    """추출된 프로토콜"""
    title: str
    description: Optional[str] = None
    materials: List[Material] = field(default_factory=list)
    equipment: List[Equipment] = field(default_factory=list)
    conditions: List[Condition] = field(default_factory=list)
    procedure: List[Step] = field(default_factory=list)
    characterization: List[Dict] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    confidence_score: float = 0.0
    extraction_timestamp: str = field(default_factory=lambda: datetime.now().isoformat())


# ==================== 프로토콜 추출기 ====================

class ProtocolExtractor:
    """프로토콜 추출 메인 클래스"""
    
    # 클래스 상수
    MAX_FILE_SIZE = 500 * 1024 * 1024  # 500MB
    MAX_TEXT_LENGTH = 500000  # 500,000자
    EXTRACTION_TIMEOUT = 60  # 60초
    
    # 섹션 패턴
    SECTION_PATTERNS = {
        'methods': r'(?i)(methods?|experimental|procedure|methodology|materials?\s+and\s+methods?)',
        'materials': r'(?i)(materials?|reagents?|chemicals?|compounds?)',
        'equipment': r'(?i)(equipment|instruments?|apparatus)',
        'procedure': r'(?i)(procedure|protocol|synthesis|preparation)',
        'characterization': r'(?i)(characterization|analysis|measurement)'
    }
    
    # 단위 매핑
    UNIT_MAPPINGS = {
        # 질량
        'g': 'g', 'gram': 'g', 'grams': 'g',
        'mg': 'mg', 'milligram': 'mg', 'milligrams': 'mg',
        'kg': 'kg', 'kilogram': 'kg', 'kilograms': 'kg',
        'μg': 'μg', 'ug': 'μg', 'microgram': 'μg',
        
        # 부피
        'l': 'L', 'liter': 'L', 'liters': 'L', 'litre': 'L',
        'ml': 'mL', 'milliliter': 'mL', 'milliliters': 'mL',
        'μl': 'μL', 'ul': 'μL', 'microliter': 'μL',
        'nl': 'nL', 'nanoliter': 'nL', 'nanoliters': 'nL',
        
        # 온도
        '°c': '°C', 'c': '°C', 'celsius': '°C',
        '°f': '°F', 'f': '°F', 'fahrenheit': '°F',
        'k': 'K', 'kelvin': 'K',
        
        # 시간
        'd': 'd', 'day': 'd', 'days': 'd',
        'h': 'h', 'hr': 'h', 'hour': 'h', 'hours': 'h',
        'min': 'min', 'minute': 'min', 'minutes': 'min',
        's': 's', 'sec': 's', 'second': 's', 'seconds': 's',
        
        # 압력
        'pa': 'Pa', 'pascal': 'Pa',
        'kpa': 'kPa', 'kilopascal': 'kPa',
        'mpa': 'MPa', 'megapascal': 'MPa',
        'gpa': 'GPa', 'gigapascal': 'GPa',
        'bar': 'bar', 'atm': 'atm', 'atmosphere': 'atm',
        'torr': 'Torr', 'mmhg': 'mmHg'

        # 퍼센트 및 비율
        '%': '%', 'percent': '%', 'percentage': '%',
        'wt%': 'wt%', 'wt.%': 'wt%', 'w/w': 'wt%', 'weight%': 'wt%',
        'v/v': 'v/v', 'vol%': 'v/v', 'volume%': 'v/v',
        'w/v': 'w/v', 'wt/vol': 'w/v',
        'mol%': 'mol%', 'mole%': 'mol%', 'at%': 'at%', 'atom%': 'at%',
    
        # 농도
        'm': 'M', 'molar': 'M', 'mol/l': 'M', 'mol/L': 'M',
        'mm': 'mM', 'millimolar': 'mM', 'mmol/l': 'mM', 'mmol/L': 'mM',
        'μm': 'μM', 'um': 'μM', 'micromolar': 'μM', 'μmol/l': 'μM',
        'nm': 'nM', 'nanomolar': 'nM', 'nmol/l': 'nM', 'nmol/L': 'nM',
        'pm': 'pM', 'picomolar': 'pM', 'pmol/l': 'pM',
    
        # 질량 농도
        'mg/ml': 'mg/mL', 'mg/mL': 'mg/mL',
        'μg/ml': 'μg/mL', 'ug/ml': 'μg/mL', 'μg/mL': 'μg/mL',
        'g/l': 'g/L', 'g/L': 'g/L',
        'mg/l': 'mg/L', 'mg/L': 'mg/L',
    
        # ppm/ppb/ppt
        'ppm': 'ppm', 'parts per million': 'ppm',
        'ppb': 'ppb', 'parts per billion': 'ppb',
        'ppt': 'ppt', 'parts per trillion': 'ppt',
    
        # 비율/배수
        'fold': 'fold', '-fold': 'fold', 'x': 'x', 'times': 'x',
        'ratio': 'ratio',
    
        # 속도
        'rpm': 'rpm', 'rev/min': 'rpm',
        'ml/min': 'mL/min', 'mL/min': 'mL/min',
        'ml/h': 'mL/h', 'mL/h': 'mL/h', 'ml/hr': 'mL/h',
        'l/min': 'L/min', 'L/min': 'L/min',
    
        # 기타 일반적인 단위
        'mesh': 'mesh',
        'kda': 'kDa', 'kilodalton': 'kDa', 'kD': 'kDa',
        'da': 'Da', 'dalton': 'Da',
        'mw': 'MW', 'molecular weight': 'MW'
    }
      

  
    
    def __init__(self):
        """초기화"""
        self.logger = logging.getLogger(__name__)
        self.error_handler = ErrorHandler()
        self._load_nlp_models()
        self._setup_patterns()
        
    def _load_nlp_models(self):
        """NLP 모델 로드"""
        try:
            # 기본 영어 모델 (작은 크기)
            self.nlp = spacy.load("en_core_web_sm")
        except:
            # 모델이 없으면 간단한 토크나이저만 사용
            self.nlp = None
            self.logger.warning("Spacy model not found. Using basic tokenization.")
            
    def _setup_patterns(self):
        """정규식 패턴 설정"""
        # 화학물질 패턴
        self.chemical_pattern = re.compile(
            r'\b([A-Z][a-z]?(\d+)?[-\(\)]?)+\b|'  # 화학식
            r'\b\d+[,-]\d+[,-]\d+\b|'  # CAS 번호
            r'\b[A-Z]{2,}[0-9]*\b'  # 약어 (예: DMSO, THF)
        )
        
        # 수량 패턴
        self.quantity_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(%|[a-zA-Zμ°]+(?:/[a-zA-Z]+)?(?:\d+)?)'
        )
        
        # 온도 패턴
        self.temperature_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*°?([CFK])\b'
        )
        
        # 시간 패턴
        self.time_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(hours?|hrs?|minutes?|mins?|seconds?|secs?|h|min|s)\b',
            re.IGNORECASE
        )

        # 농도 패턴 추가
        self.concentration_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(M|mM|μM|uM|nM|pM)\b'
        )
    
        # 퍼센트 패턴 추가
        self.percent_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(wt\.?%|w/w|v/v|w/v|mol%|at%|%)'
        )
    
        # ppm/ppb/ppt 패턴 추가
        self.ppm_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(ppm|ppb|ppt)\b'
        )
    
        # 질량 농도 패턴 추가
        self.mass_concentration_pattern = re.compile(
            r'(\d+(?:\.\d+)?)\s*(mg/mL|μg/mL|ug/mL|g/L|mg/L)'
        )

    # ==================== 파일 처리 ====================
    
    def extract_from_file(
        self, 
        file_path: Union[str, Path],
        file_type: Optional[str] = None
    ) -> Protocol:
        """파일에서 프로토콜 추출"""
        file_path = Path(file_path)
        
        # 파일 검증
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if file_path.stat().st_size > self.MAX_FILE_SIZE:
            self.error_handler.handle_error(
                Exception("File too large"),
                context={'file_size': file_path.stat().st_size, 'max_size': self.MAX_FILE_SIZE},
                error_code='4203'
            )
            
        # 파일 타입 결정
        if not file_type:
            file_type = DocumentType.from_extension(file_path.suffix)
            
        if not file_type:
            self.error_handler.handle_error(
                Exception("Unsupported file type"),
                context={'file_type': file_path.suffix},
                error_code='4200'
            )
            
        # 타임아웃 처리
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(self._extract_with_type, file_path, file_type)
            try:
                return future.result(timeout=self.EXTRACTION_TIMEOUT)
            except FuturesTimeoutError:
                self.error_handler.handle_error(
                    Exception("Extraction timeout"),
                    context={'file': str(file_path)},
                    error_code='4208'
                )
                
    def _extract_with_type(self, file_path: Path, file_type: DocumentType) -> Protocol:
        """파일 타입별 추출"""
        extractors = {
            DocumentType.PDF: self._extract_from_pdf,
            DocumentType.DOCX: self._extract_from_docx,
            DocumentType.TXT: self._extract_from_txt,
            DocumentType.HTML: self._extract_from_html,
            DocumentType.MD: self._extract_from_markdown,
            DocumentType.RTF: self._extract_from_rtf
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"No extractor for type: {file_type}")
            
        # 텍스트 추출
        text = extractor(file_path)
        
        # 텍스트 길이 체크
        if len(text) > self.MAX_TEXT_LENGTH:
            self.logger.warning(f"Text truncated from {len(text)} to {self.MAX_TEXT_LENGTH} characters")
            text = text[:self.MAX_TEXT_LENGTH]
            
        # 프로토콜 추출
        return self._extract_protocol_from_text(text, file_path.name)
        
    # ==================== 파일별 추출 메서드 ====================
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """PDF에서 텍스트 추출"""
        text_parts = []
        
        # PyPDF2 시도
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text_parts.append(page.extract_text())
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {page_num}: {e}")
        except Exception as e:
            self.logger.warning(f"PyPDF2 failed: {e}")
            
        # pdfplumber로 재시도 (테이블 추출에 더 좋음)
        if not text_parts:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                            
                        # 테이블 추출
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                # 테이블을 텍스트로 변환
                                table_text = self._table_to_text(table)
                                text_parts.append(table_text)
            except Exception as e:
                self.logger.warning(f"pdfplumber failed: {e}")
                
        # OCR 시도
        if not text_parts and OCR_AVAILABLE:
            text_parts = self._ocr_pdf(file_path)
            
        if not text_parts:
            self.error_handler.handle_error(
                Exception("No text extracted from PDF"),
                context={'file': str(file_path)},
                error_code='4202'
            )
            
        return '\n\n'.join(text_parts)
        
    def _extract_from_docx(self, file_path: Path) -> str:
        """Word 문서에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 단락 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # 테이블 추출
            for table in doc.tables:
                table_text = self._extract_table_from_docx(table)
                if table_text:
                    text_parts.append(table_text)
                    
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.error_handler.handle_error(
                Exception(f"Error reading DOCX: {e}"),
                context={'file': str(file_path)},
                error_code='4204'
            )
            
    def _extract_from_txt(self, file_path: Path) -> str:
        """텍스트 파일 읽기 (인코딩 자동 감지)"""
        # 먼저 인코딩 감지
        encoding = self.detect_encoding(file_path)
        
        if encoding:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except Exception as e:
                self.logger.warning(f"Failed with detected encoding {encoding}: {e}")
                
        # 여러 인코딩 시도
        return self.try_multiple_encodings(
            file_path,
            'utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'cp949', 'gbk'
        )
        
    def _extract_from_html(self, file_path: Path) -> str:
        """HTML에서 텍스트 추출"""
        encoding = self.detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding or 'utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                
            # 스크립트와 스타일 제거
            for script in soup(["script", "style"]):
                script.decompose()
                
            # 텍스트 추출
            text = soup.get_text()
            
            # 여러 공백을 하나로
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            self.error_handler.handle_error(
                Exception(f"Error reading HTML: {e}"),
                context={'file': str(file_path)},
                error_code='4204'
            )
            
    def _extract_from_markdown(self, file_path: Path) -> str:
        """Markdown에서 텍스트 추출"""
        text = self._extract_from_txt(file_path)
        
        # Markdown을 HTML로 변환 후 텍스트 추출
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        
        return soup.get_text()
        
    def _extract_from_rtf(self, file_path: Path) -> str:
        """RTF에서 텍스트 추출"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
            
        return rtf_to_text(rtf_content)
        
    # ==================== 프로토콜 추출 ====================
    
    def _extract_protocol_from_text(self, text: str, source_name: str = "Unknown") -> Protocol:
        """텍스트에서 프로토콜 추출"""
        # 섹션 분리
        sections = self._identify_sections(text)
        
        if not sections:
            self.error_handler.handle_error(
                Exception("No protocol sections found"),
                context={'source': source_name},
                error_code='4202'
            )
            
        # 프로토콜 객체 생성
        protocol = Protocol(
            title=self._extract_title(text, source_name),
            description=self._extract_description(sections)
        )
        
        # 각 요소 추출
        protocol.materials = self._extract_materials(sections)
        protocol.equipment = self._extract_equipment(sections)
        protocol.conditions = self._extract_conditions(sections)
        protocol.procedure = self._extract_procedure(sections)
        protocol.characterization = self._extract_characterization(sections)
        
        # 메타데이터
        protocol.metadata = {
            'source': source_name,
            'sections_found': list(sections.keys()),
            'text_length': len(text),
            'extraction_method': 'automatic'
        }
        
        # 신뢰도 점수 계산
        protocol.confidence_score = self._calculate_confidence(protocol)
        
        # 최소 요구사항 체크
        if not protocol.materials and not protocol.procedure:
            self.error_handler.handle_error(
                Exception("Insufficient protocol data"),
                context={'source': source_name},
                error_code='4209'
            )
            
        return protocol
        
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """텍스트에서 섹션 식별"""
        sections = {}
        lines = text.split('\n')
        
        current_section = None
        current_content = []
        
        for line in lines:
            # 섹션 헤더 찾기
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.search(pattern, line):
                    # 이전 섹션 저장
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                        
                    current_section = section_name
                    current_content = []
                    break
            else:
                # 현재 섹션에 내용 추가
                if current_section:
                    current_content.append(line)
                    
        # 마지막 섹션 저장
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
            
        # 전체 텍스트에서 패턴 매칭 (섹션이 명확하지 않은 경우)
        if not sections:
            sections['full_text'] = text
            
        return sections
        
    def _extract_title(self, text: str, source_name: str) -> str:
        """제목 추출"""
        # 첫 몇 줄에서 제목 찾기
        lines = text.split('\n')[:10]
        
        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 200:
                # 제목일 가능성이 높은 패턴
                if not any(char in line for char in ['.', ':', ';']) or line.endswith(':'):
                    return line.rstrip(':')
                    
        return f"Protocol from {source_name}"
        
    def _extract_description(self, sections: Dict[str, str]) -> Optional[str]:
        """설명 추출"""
        # Abstract나 Introduction에서 추출
        for section in ['abstract', 'introduction', 'full_text']:
            if section in sections:
                text = sections[section][:500]  # 처음 500자
                # 첫 문단 추출
                paragraphs = text.split('\n\n')
                if paragraphs:
                    return paragraphs[0].strip()
                    
        return None
        
    def _extract_materials(self, sections: Dict[str, str]) -> List[Material]:
        """재료 추출"""
        materials = []
        material_texts = []
        
        # 재료 섹션 찾기
        for section_name, content in sections.items():
            if 'material' in section_name or section_name == 'full_text':
                material_texts.append(content)
                
        # 재료 추출
        for text in material_texts:
            # 화학물질 패턴 매칭
            chemicals = self.chemical_pattern.findall(text)
            
            # 각 화학물질에 대해
            for chemical in set(chemicals):
                if isinstance(chemical, tuple):
                    chemical = chemical[0]
                    
                if len(chemical) < 2:  # 너무 짧은 것 제외
                    continue
                    
                material = Material(name=chemical)
                
                # CAS 번호 찾기
                cas_match = re.search(rf'{re.escape(chemical)}[^0-9]*(\d{{1,7}}-\d{{2}}-\d{{1}})', text)
                if cas_match:
                    material.cas_number = cas_match.group(1)
                    
                # 공급업체 찾기
                supplier_patterns = [
                    rf'{re.escape(chemical)}[^,;]*(?:from|purchased from|supplied by)\s+([A-Z][A-Za-z\s&-]+)',
                    rf'([A-Z][A-Za-z\s&-]+)[,\s]+{re.escape(chemical)}'
                ]
                
                for pattern in supplier_patterns:
                    supplier_match = re.search(pattern, text)
                    if supplier_match:
                        material.supplier = supplier_match.group(1).strip()
                        break
                        
                # 순도 찾기
                purity_patterns = [
                    rf'{re.escape(chemical)}[^,;]*(\d+(?:\.\d+)?%)',
                    rf'{re.escape(chemical)}[^,;]*\((\d+(?:\.\d+)?%)\)',
                    rf'{re.escape(chemical)}[^,;]*purity[:\s]*(\d+(?:\.\d+)?%)',
                    rf'{re.escape(chemical)}[^,;]*\(>(\d+(?:\.\d+)?%)\)',
                    rf'{re.escape(chemical)}[^,;]*(≥\s*\d+(?:\.\d+)?%)'
                ]
    
                for pattern in purity_patterns:
                    purity_match = re.search(pattern, text, re.IGNORECASE)
                    if purity_match:
                        material.purity = purity_match.group(1)
                        break
                    
                # 양 찾기
                quantity_match = re.search(rf'(\d+(?:\.\d+)?)\s*([a-zA-Z]+)\s+(?:of\s+)?{re.escape(chemical)}', text)
                if quantity_match:
                    material.amount = quantity_match.group(1)
                    material.unit = self._standardize_unit(quantity_match.group(2))
                    
                materials.append(material)
                
        # NLP로 추가 추출 (spacy 사용 가능한 경우)
        if self.nlp and materials:
            materials = self._enhance_materials_with_nlp(materials, material_texts)
            
        return materials
        
    def _extract_equipment(self, sections: Dict[str, str]) -> List[Equipment]:
        """장비 추출"""
        equipment_list = []
        
        # 장비 패턴
        equipment_patterns = [
            r'([A-Z][A-Za-z\s-]+(?:meter|scope|analyzer|system|instrument|apparatus|device))',
            r'([A-Z]{2,}(?:-\d+)?)',  # 모델명 (예: UV-2600)
            r'(?:using|with|by)\s+(?:a|an|the)?\s*([A-Z][A-Za-z\s-]+)'
        ]
        
        for section_name, content in sections.items():
            for pattern in equipment_patterns:
                matches = re.findall(pattern, content)
                
                for match in matches:
                    if len(match) > 3:  # 너무 짧은 것 제외
                        equipment = Equipment(name=match.strip())
                        
                        # 제조사 찾기
                        manufacturer_match = re.search(
                            rf'{re.escape(match)}[^,;]*\(([A-Z][A-Za-z\s&.-]+)\)',
                            content
                        )
                        if manufacturer_match:
                            equipment.manufacturer = manufacturer_match.group(1)
                            
                        # 모델 찾기
                        model_match = re.search(
                            rf'{re.escape(match)}[^,;]*(?:model|Model)\s*([A-Z0-9-]+)',
                            content
                        )
                        if model_match:
                            equipment.model = model_match.group(1)
                            
                        equipment_list.append(equipment)
                        
        # 중복 제거
        seen = set()
        unique_equipment = []
        for eq in equipment_list:
            key = eq.name.lower()
            if key not in seen:
                seen.add(key)
                unique_equipment.append(eq)
                
        return unique_equipment
        
    def _extract_conditions(self, sections: Dict[str, str]) -> List[Condition]:
        """실험 조건 추출"""
        conditions = []
        
        # 모든 섹션에서 조건 찾기
        for section_name, content in sections.items():
            # 온도
            temp_matches = self.temperature_pattern.findall(content)
            for value, unit in temp_matches:
                conditions.append(Condition(
                    parameter="temperature",
                    value=float(value),
                    unit=f"°{unit}" if unit in ['C', 'F'] else unit
                ))
                
            # 시간
            time_matches = self.time_pattern.findall(content)
            for value, unit in time_matches:
                conditions.append(Condition(
                    parameter="time",
                    value=float(value),
                    unit=self._standardize_unit(unit)
                ))
                
            # 압력
            pressure_patterns = [
                r'(\d+(?:\.\d+)?)\s*(atm|bar|Pa|kPa|MPa|Torr|mmHg)',
                r'(?:at|under)\s+(\d+(?:\.\d+)?)\s*(atm|bar|Pa|kPa|MPa|Torr|mmHg)'
            ]
            
            for pattern in pressure_patterns:
                pressure_matches = re.findall(pattern, content, re.IGNORECASE)
                for value, unit in pressure_matches:
                    conditions.append(Condition(
                        parameter="pressure",
                        value=float(value),
                        unit=self._standardize_unit(unit)
                    ))
                    
            # pH
            ph_matches = re.findall(r'pH\s*(?:of|=)?\s*(\d+(?:\.\d+)?)', content)
            for value in ph_matches:
                conditions.append(Condition(
                    parameter="pH",
                    value=float(value)
                ))
                
            # 회전 속도
            rpm_matches = re.findall(r'(\d+)\s*rpm', content, re.IGNORECASE)
            for value in rpm_matches:
                conditions.append(Condition(
                    parameter="rotation_speed",
                    value=int(value),
                    unit="rpm"
                ))

            # 농도 조건 추가
            conc_matches = self.concentration_pattern.findall(content)
            for value, unit in conc_matches:
                conditions.append(Condition(
                    parameter="concentration",
                    value=float(value),
                    unit=self._standardize_unit(unit)
                ))
        
            # 퍼센트 조건 추가
            percent_matches = self.percent_pattern.findall(content)
            for value, unit in percent_matches:
                # 퍼센트 타입 구분
                param_name = "percentage"
                if 'wt' in unit or 'w/w' in unit:
                    param_name = "weight_percentage"
                elif 'v/v' in unit:
                    param_name = "volume_percentage"
                elif 'mol' in unit:
                    param_name = "molar_percentage"
                
                conditions.append(Condition(
                    parameter=param_name,
                    value=float(value),
                    unit=self._standardize_unit(unit)
                ))
        
            # 유속/속도 추가
            flow_patterns = [
                r'(\d+(?:\.\d+)?)\s*(mL/min|L/min|mL/h|μL/min)',
                r'flow\s+rate[:\s]*(\d+(?:\.\d+)?)\s*(mL/min|L/min)'
            ]
        
            for pattern in flow_patterns:
                flow_matches = re.findall(pattern, content, re.IGNORECASE)
                for value, unit in flow_matches:
                    conditions.append(Condition(
                        parameter="flow_rate",
                        value=float(value),
                        unit=self._standardize_unit(unit)
                    ))
      
        return conditions
        
    def _extract_procedure(self, sections: Dict[str, str]) -> List[Step]:
        """실험 절차 추출"""
        steps = []
        procedure_text = ""
        
        # 절차 섹션 찾기
        for section_name, content in sections.items():
            if 'procedure' in section_name or 'method' in section_name:
                procedure_text += content + "\n"
                
        if not procedure_text and 'full_text' in sections:
            procedure_text = sections['full_text']
            
        # 단계 분리 패턴
        step_patterns = [
            r'(?:Step\s+)?(\d+)[.):]\s*([^.]+(?:\.[^.]+)*?)(?=(?:Step\s+)?\d+[.):]\s*|\Z)',
            r'(?:^|\n)([A-Z][^.]+(?:\.[^.]+)*?)(?=\n[A-Z]|\Z)',
            r'(?:First|Then|Next|Finally|Subsequently)[,\s]+([^.]+\.)'
        ]
        
        step_number = 1
        
        for pattern in step_patterns:
            matches = re.findall(pattern, procedure_text, re.MULTILINE | re.DOTALL)
            
            if matches:
                for match in matches:
                    if isinstance(match, tuple):
                        if len(match) == 2 and match[0].isdigit():
                            step_text = match[1].strip()
                        else:
                            step_text = match[0].strip()
                    else:
                        step_text = match.strip()
                        
                    if len(step_text) > 10:  # 의미 있는 길이
                        step = Step(
                            number=step_number,
                            action=step_text
                        )
                        
                        # 온도 추출
                        temp_match = self.temperature_pattern.search(step_text)
                        if temp_match:
                            step.temperature = float(temp_match.group(1))
                            
                        # 시간 추출
                        time_match = self.time_pattern.search(step_text)
                        if time_match:
                            step.duration = f"{time_match.group(1)} {time_match.group(2)}"
                            
                        # 중요 단계 표시
                        if any(word in step_text.lower() for word in ['critical', 'important', 'caution', 'warning']):
                            step.critical = True
                            
                        steps.append(step)
                        step_number += 1
                        
                break  # 첫 번째 성공한 패턴만 사용
                
        # 단계가 없으면 문장 기반 분리
        if not steps:
            sentences = re.split(r'[.!?]+', procedure_text)
            for i, sentence in enumerate(sentences):
                sentence = sentence.strip()
                if len(sentence) > 20 and any(verb in sentence.lower() for verb in 
                    ['add', 'mix', 'heat', 'cool', 'stir', 'dissolve', 'filter', 'wash', 'dry']):
                    steps.append(Step(
                        number=i + 1,
                        action=sentence
                    ))
                    
        return steps
        
    def _extract_characterization(self, sections: Dict[str, str]) -> List[Dict]:
        """특성 분석 방법 추출"""
        characterization = []
        
        # 분석 방법 패턴
        analysis_methods = {
            'NMR': r'(?:^|H\s+|C\s+)?NMR',
            'FTIR': r'FT-?IR',
            'UV-Vis': r'UV-?Vis',
            'GPC': r'GPC|SEC',
            'DSC': r'DSC',
            'TGA': r'TGA',
            'XRD': r'XRD',
            'SEM': r'SEM',
            'TEM': r'TEM',
            'AFM': r'AFM',
            'XPS': r'XPS',
            'MS': r'(?:mass\s+spec|MS)',
            'GC': r'GC(?:-MS)?',
            'HPLC': r'HPLC'
        }
        
        for section_name, content in sections.items():
            for method_name, pattern in analysis_methods.items():
                if re.search(pattern, content, re.IGNORECASE):
                    method_info = {
                        'method': method_name,
                        'parameters': {}
                    }
                    
                    # 파라미터 추출 (예: 주파수, 온도 등)
                    param_match = re.search(
                        rf'{pattern}[^.;]*?(\d+)\s*(MHz|Hz|nm|°C)',
                        content,
                        re.IGNORECASE
                    )
                    if param_match:
                        method_info['parameters']['value'] = param_match.group(1)
                        method_info['parameters']['unit'] = param_match.group(2)
                        
                    characterization.append(method_info)
                    
        return characterization
        
    # ==================== 헬퍼 메서드 ====================
    
    def _standardize_unit(self, unit: str) -> str:
        """단위 표준화"""
        unit_lower = unit.lower()
        return self.UNIT_MAPPINGS.get(unit_lower, unit)
        
    def _table_to_text(self, table: List[List]) -> str:
        """테이블을 텍스트로 변환"""
        text_parts = []
        
        for row in table:
            if row:
                # None 값 제거하고 문자열로 변환
                row_text = ' | '.join(str(cell) for cell in row if cell is not None)
                text_parts.append(row_text)
                
        return '\n'.join(text_parts)
        
    def _extract_table_from_docx(self, table) -> str:
        """Word 테이블에서 텍스트 추출"""
        text_parts = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(' | '.join(row_text))
                
        return '\n'.join(text_parts)
        
    def _calculate_confidence(self, protocol: Protocol) -> float:
        """신뢰도 점수 계산"""
        score = 0.0
        weights = {
            'materials': 0.3,
            'procedure': 0.3,
            'conditions': 0.2,
            'equipment': 0.1,
            'characterization': 0.1
        }
        
        if protocol.materials:
            score += weights['materials'] * min(len(protocol.materials) / 5, 1.0)
            
        if protocol.procedure:
            score += weights['procedure'] * min(len(protocol.procedure) / 5, 1.0)
            
        if protocol.conditions:
            score += weights['conditions'] * min(len(protocol.conditions) / 3, 1.0)
            
        if protocol.equipment:
            score += weights['equipment'] * min(len(protocol.equipment) / 2, 1.0)
            
        if protocol.characterization:
            score += weights['characterization'] * min(len(protocol.characterization) / 2, 1.0)
            
        return round(score, 2)
        
    def _enhance_materials_with_nlp(
        self, 
        materials: List[Material], 
        texts: List[str]
    ) -> List[Material]:
        """NLP를 사용한 재료 정보 향상"""
        if not self.nlp:
            return materials
            
        # 텍스트 분석
        for text in texts:
            doc = self.nlp(text[:1000000])  # spacy 제한
            
            # 명명 개체 인식
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'PRODUCT']:
                    # 공급업체일 가능성
                    for material in materials:
                        if material.name in ent.sent.text and not material.supplier:
                            material.supplier = ent.text
                            
        return materials
        
    # ==================== OCR 처리 ====================
    
    def _ocr_pdf(self, file_path: Path) -> List[str]:
        """PDF OCR 처리"""
        if not OCR_AVAILABLE:
            self.logger.warning("OCR not available. Install pytesseract and PIL.")
            return []
            
        text_parts = []
        
        try:
            # PDF를 이미지로 변환 (pdf2image 필요)
            from pdf2image import convert_from_path
            
            images = convert_from_path(file_path, dpi=300)
            
            for i, image in enumerate(images):
                # 이미지 향상
                enhanced_image = self.enhance_image(image)
                
                # OCR 수행
                text = pytesseract.image_to_string(enhanced_image)
                if text.strip():
                    text_parts.append(text)
                    
        except Exception as e:
            self.logger.warning(f"OCR failed: {e}")
            
        return text_parts
        
    # ==================== 에러 복구 메서드 ====================
    
    @staticmethod
    def try_multiple_encodings(filepath: Union[str, Path], *encodings, **context) -> str:
        """여러 인코딩으로 파일 읽기 시도"""
        filepath = Path(filepath)
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
                
        # 모든 인코딩 실패 시 바이너리로 읽고 디코드 시도
        with open(filepath, 'rb') as f:
            content = f.read()
            
        # 손실 있는 디코딩
        return content.decode('utf-8', errors='replace')
        
    @staticmethod
    def detect_encoding(filepath: Union[str, Path], **context) -> Optional[str]:
        """파일 인코딩 자동 감지"""
        filepath = Path(filepath)
        
        try:
            with open(filepath, 'rb') as f:
                raw_data = f.read(10000)  # 처음 10KB만 읽기
                
            result = chardet.detect(raw_data)
            
            if result['confidence'] > 0.7:
                return result['encoding']
                
        except Exception:
            pass
            
        return None
        
    @staticmethod
    def enhance_image(image: Image.Image, *operations, **context) -> Image.Image:
        """이미지 품질 개선"""
        from PIL import ImageEnhance, ImageFilter
        
        if not operations:
            operations = ['contrast', 'sharpness']
            
        for operation in operations:
            if operation == 'contrast':
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(1.5)
            elif operation == 'sharpness':
                enhancer = ImageEnhance.Sharpness(image)
                image = enhancer.enhance(2.0)
            elif operation == 'denoise':
                image = image.filter(ImageFilter.MedianFilter(size=3))
                
        return image
        
    # ==================== 다중 파일 처리 ====================
    
    def extract_from_multiple_files(
        self, 
        file_paths: List[Union[str, Path]],
        merge: bool = True
    ) -> Union[Protocol, List[Protocol]]:
        """여러 파일에서 프로토콜 추출"""
        protocols = []
        failed_files = []
        
        for file_path in file_paths:
            try:
                protocol = self.extract_from_file(file_path)
                protocols.append(protocol)
            except Exception as e:
                self.logger.warning(f"Failed to extract from {file_path}: {e}")
                failed_files.append(str(file_path))
                
        if failed_files:
            self.error_handler.handle_error(
                Exception("Some files failed"),
                context={
                    'failed_count': len(failed_files),
                    'total_count': len(file_paths),
                    'failed_files': failed_files
                },
                error_code='4206'
            )
            
        if not protocols:
            raise ValueError("No protocols extracted from any file")
            
        if merge and len(protocols) > 1:
            return self._merge_protocols(protocols)
            
        return protocols[0] if len(protocols) == 1 else protocols
        
    def _merge_protocols(self, protocols: List[Protocol]) -> Protocol:
        """여러 프로토콜 병합"""
        merged = Protocol(
            title=f"Merged Protocol from {len(protocols)} sources",
            description="Combined protocol from multiple sources"
        )
        
        # 재료 병합 (중복 제거)
        material_dict = {}
        for protocol in protocols:
            for material in protocol.materials:
                key = material.name.lower()
                if key not in material_dict:
                    material_dict[key] = material
                    
        merged.materials = list(material_dict.values())
        
        # 장비 병합
        equipment_dict = {}
        for protocol in protocols:
            for equipment in protocol.equipment:
                key = equipment.name.lower()
                if key not in equipment_dict:
                    equipment_dict[key] = equipment
                    
        merged.equipment = list(equipment_dict.values())
        
        # 조건 병합
        for protocol in protocols:
            merged.conditions.extend(protocol.conditions)
            
        # 절차 병합 (순서대로)
        step_number = 1
        for protocol in protocols:
            for step in protocol.procedure:
                new_step = Step(
                    number=step_number,
                    action=step.action,
                    duration=step.duration,
                    temperature=step.temperature,
                    conditions=step.conditions,
                    notes=step.notes,
                    critical=step.critical
                )
                merged.procedure.append(new_step)
                step_number += 1
                
        # 특성 분석 병합
        char_dict = {}
        for protocol in protocols:
            for char in protocol.characterization:
                key = char.get('method', '').lower()
                if key not in char_dict:
                    char_dict[key] = char
                    
        merged.characterization = list(char_dict.values())
        
        # 메타데이터
        merged.metadata = {
            'sources': [p.metadata.get('source', 'Unknown') for p in protocols],
            'merge_count': len(protocols),
            'merge_timestamp': datetime.now().isoformat()
        }
        
        # 신뢰도는 평균
        if protocols:
            merged.confidence_score = sum(p.confidence_score for p in protocols) / len(protocols)
            
        return merged
        
    # ==================== 내보내기 ====================
    
    def to_json(self, protocol: Protocol, indent: int = 2) -> str:
        """프로토콜을 JSON으로 변환"""
        return json.dumps(asdict(protocol), indent=indent, ensure_ascii=False)
        
    def to_template(self, protocol: Protocol) -> Dict[str, Any]:
        """프로토콜을 실험 템플릿으로 변환"""
        template = {
            'name': protocol.title,
            'description': protocol.description,
            'factors': [],
            'responses': [],
            'constraints': []
        }
        
        # 조건을 요인으로 변환
        for condition in protocol.conditions:
            factor = {
                'name': condition.parameter,
                'type': 'continuous',
                'unit': condition.unit
            }
            
            if condition.value:
                factor['default'] = condition.value
                
            template['factors'].append(factor)
            
        # 특성 분석을 반응변수로
        for char in protocol.characterization:
            template['responses'].append({
                'name': char.get('method', 'Unknown'),
                'type': 'continuous'
            })
            
        return template
        
    # ==================== URL 처리 ====================
    
    def extract_from_url(self, url: str) -> Protocol:
        """URL에서 프로토콜 추출"""
        import requests
        
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp:
                tmp.write(response.content)
                tmp_path = tmp.name
                
            try:
                # HTML로 처리
                return self.extract_from_file(tmp_path, DocumentType.HTML)
            finally:
                # 임시 파일 삭제
                os.unlink(tmp_path)
                
        except Exception as e:
            self.error_handler.handle_error(
                Exception(f"Failed to fetch URL: {e}"),
                context={'url': url},
                error_code='4207'
            )


# ==================== 편의 함수 ====================

def extract_protocol(
    source: Union[str, Path, List[Union[str, Path]]],
    merge: bool = True
) -> Union[Protocol, List[Protocol]]:
    """프로토콜 추출 편의 함수"""
    extractor = ProtocolExtractor()
    
    if isinstance(source, list):
        return extractor.extract_from_multiple_files(source, merge=merge)
    elif str(source).startswith(('http://', 'https://')):
        return extractor.extract_from_url(str(source))
    else:
        return extractor.extract_from_file(source)


def protocol_to_json(protocol: Protocol, filepath: Optional[Union[str, Path]] = None) -> str:
    """프로토콜을 JSON으로 저장"""
    extractor = ProtocolExtractor()
    json_str = extractor.to_json(protocol)
    
    if filepath:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(json_str)
            
    return json_str
